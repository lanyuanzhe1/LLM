"""Tests for app.ingest.reader — format-routing text extraction module."""

from pathlib import Path

import fitz
import pytest
from docx import Document

from app.ingest.reader import SUPPORTED_EXTENSIONS, read_file


class TestSupportedExtensions:
    def test_supported_extensions_are_v1_scope(self):
        assert SUPPORTED_EXTENSIONS == frozenset({".pdf", ".docx", ".txt", ".md"})


class TestTextFiles:
    def test_txt_returns_full_content(self, tmp_path):
        file = tmp_path / "test.txt"
        file.write_text("低温储粮技术", encoding="utf-8")
        result = read_file(file)
        assert result == "低温储粮技术"

    def test_md_returns_full_content(self, tmp_path):
        file = tmp_path / "readme.md"
        file.write_text("# 标题\n内容", encoding="utf-8")
        result = read_file(file)
        assert "内容" in result

    def test_empty_file_returns_empty(self, tmp_path):
        file = tmp_path / "empty.txt"
        file.write_text("", encoding="utf-8")
        result = read_file(file)
        assert result == ""

    def test_missing_file_returns_empty(self, tmp_path):
        nonexistent = tmp_path / "does_not_exist.pdf"
        result = read_file(nonexistent)
        assert result == ""


class TestDocx:
    def test_docx_extracts_paragraphs(self, tmp_path):
        file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("粮食储藏技术规范")
        doc.add_paragraph("害虫防治方法")
        doc.save(str(file))
        result = read_file(file)
        assert "粮食储藏技术规范" in result
        assert "害虫防治方法" in result


class TestPdf:
    def test_pdf_extracts_text(self, tmp_path):
        file = tmp_path / "test.pdf"
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        page.insert_text((72, 72), "Grain Storage Pest Control")
        pdf_doc.save(str(file))
        pdf_doc.close()
        result = read_file(file)
        assert "Grain Storage Pest Control" in result

    def test_pdf_with_no_text_returns_empty(self, tmp_path):
        file = tmp_path / "blank.pdf"
        pdf_doc = fitz.open()
        pdf_doc.new_page()  # blank page, no text inserted
        pdf_doc.save(str(file))
        pdf_doc.close()
        result = read_file(file)
        assert result == ""


class TestUnsupported:
    def test_unsupported_extension_returns_empty(self, tmp_path):
        file = tmp_path / "slides.pptx"
        file.write_text("fake content", encoding="utf-8")
        result = read_file(file)
        assert result == ""
